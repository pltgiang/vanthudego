import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_header(doc):
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Simplified header since we can't easily draw the exact logo
    run = p.add_run("DEGO HOLDING\n")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 174, 239)
    
    run2 = p.add_run("B19, ĐDC Cần Thơ, Khu DCVH Tây Đô, P.Hưng Thạnh, Q.Cái Răng, TP. Cần Thơ\n")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(128, 128, 128)
    
    run3 = p.add_run("0782.99.27.11 - 0898.00.11.13 | hr.degoholding@gmail.com\n")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(128, 128, 128)
    
    run4 = p.add_run("www.degoholding.com")
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add a line
    p_border = doc.add_paragraph()
    p_border.add_run("_"*80).font.color.rgb = RGBColor(200, 200, 200)
    p_border.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_leave_request():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    add_header(doc)
    
    doc.add_paragraph("\n")
    title = doc.add_paragraph("ĐƠN XIN NGHỈ PHÉP")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph("\n")
    p = doc.add_paragraph("Kính gửi:\t- Trưởng phòng/Bộ phận/Nhóm: ")
    p.runs[0].bold = True
    p.add_run("{department}")
    p = doc.add_paragraph("\t\t- Trưởng phòng HCNS.")
    p.runs[0].bold = True
    
    doc.add_paragraph("\n")
    # Table for info to align well
    table = doc.add_table(rows=2, cols=2)
    cell = table.cell(0,0)
    cell.text = "Tôi tên là: {fullname}"
    cell = table.cell(0,1)
    cell.text = "SĐT: {phone}"
    
    cell = table.cell(1,0)
    cell.text = "Chức vụ: {position}"
    cell = table.cell(1,1)
    cell.text = "Phòng: {department}"
    
    doc.add_paragraph("\nNay tôi làm đơn này xin phép Trưởng Bộ phận cho tôi được nghỉ phép ngày {leave_date}.")
    doc.add_paragraph("Lý do: {reason}")
    doc.add_paragraph("Nghỉ theo diện: {c_kl} Không lương   {c_pn} Phép năm   {c_no} Nghỉ ốm   {c_cd} Nghỉ chế độ")
    doc.add_paragraph("Người quản lý thay công việc: {backup_person}")
    doc.add_paragraph("Rất mong Trưởng Bộ phận công ty xem xét và chấp thuận.")
    
    p_date = doc.add_paragraph("\nCần Thơ, ngày {day} tháng {month} năm {year}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.runs[0].italic = True
    
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.cell(0,0).text = "P.HCNS"
    sig_table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_table.cell(0,0).paragraphs[0].runs[0].bold = True
    
    sig_table.cell(0,1).text = "Trưởng Phòng/Bộ phận"
    sig_table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_table.cell(0,1).paragraphs[0].runs[0].bold = True
    
    sig_table.cell(0,2).text = "Người làm đơn"
    sig_table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_table.cell(0,2).paragraphs[0].runs[0].bold = True
    
    # Padding for signature
    sig_table.cell(1,0).text = "\n\n\n{hcns_name}"
    sig_table.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_table.cell(1,0).paragraphs[0].runs[0].bold = True
    
    sig_table.cell(1,1).text = "\n\n\n{manager_name}"
    sig_table.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_table.cell(1,1).paragraphs[0].runs[0].bold = True
    
    sig_table.cell(1,2).text = "\n\n\n{fullname}"
    sig_table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_table.cell(1,2).paragraphs[0].runs[0].bold = True
    
    os.makedirs('frontend/public/templates', exist_ok=True)
    doc.save('frontend/public/templates/Don_Xin_Nghi_Phep.docx')

def create_handover():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    add_header(doc)
    
    t = doc.add_table(rows=1, cols=2)
    t.cell(0,0).text = "Số: {doc_no}/2026/BBBG - DEGO"
    
    p = t.cell(0,1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Cần Thơ, ngày {day} tháng {month} năm {year}").italic = True
    
    doc.add_paragraph("\n")
    title = doc.add_paragraph("BIÊN BẢN GIAO NHẬN")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph("\nHôm nay, ngày {day} tháng {month} năm {year}, tại văn phòng làm việc CÔNG TY TNHH DEGO HOLDING (Địa chỉ: B19, ĐDC Cần Thơ, Khu DCVH Tây Đô, P.Hưng Thạnh, Q.Cái Răng, TP. Cần Thơ), chúng tôi gồm:\n")
    
    p = doc.add_paragraph()
    p.add_run("BÊN A (Bên bàn giao): CÔNG TY TNHH DEGO HOLDING").bold = True
    doc.add_paragraph("- Địa chỉ: B19 đường dẫn cầu Cần Thơ, QL 1A, Khu DC TTVH Tây Đô, Phường Hưng Thạnh, Quận Cái Răng, TP. Cần Thơ.")
    doc.add_paragraph("- Mã số thuế: 1801722464")
    p = doc.add_paragraph("- Đại diện Ông/bà: ")
    p.add_run("{a_name}").bold = True
    doc.add_paragraph("- Chức vụ: {a_position}")
    doc.add_paragraph("- Điện thoại: {a_phone}")
    
    doc.add_paragraph("\n")
    
    p = doc.add_paragraph()
    p.add_run("BÊN B (Bên nhận): CÔNG TY TNHH DEGO HOLDING").bold = True
    doc.add_paragraph("- Địa chỉ: B19 đường dẫn cầu Cần Thơ, QL 1A, Khu DC TTVH Tây Đô, Phường Hưng Thạnh, Quận Cái Răng, TP. Cần Thơ.")
    doc.add_paragraph("- Mã số thuế: 1801722464")
    p = doc.add_paragraph("- Đại diện Ông/bà: ")
    p.add_run("{b_name}").bold = True
    doc.add_paragraph("- Chức vụ: {b_position}")
    doc.add_paragraph("- Điện thoại: {b_phone}")
    
    doc.add_paragraph("\n(Nội dung bàn giao...)\n")
    
    os.makedirs('frontend/public/templates', exist_ok=True)
    doc.save('frontend/public/templates/Bien_Ban_Ban_Giao.docx')

if __name__ == "__main__":
    create_leave_request()
    create_handover()
    print("Templates generated successfully.")
